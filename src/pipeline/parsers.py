from __future__ import annotations

import csv
from pathlib import Path

from src.pipeline.document import Cell, Document, Line


def parse_txt(path: Path) -> Document:
    with open(path, encoding="utf-8") as f:
        raw_lines = f.readlines()
    lines = [
        Line(cells=[Cell(text=text.rstrip("\n\r"))], line_number=i + 1)
        for i, text in enumerate(raw_lines)
    ]
    return Document(lines=lines, source_path=path, source_format="txt")


def parse_csv(path: Path) -> Document:
    lines: list[Line] = []
    headers: list[str] | None = None
    with open(path, newline="", encoding="utf-8") as f:
        reader = csv.reader(f)
        for row_idx, row in enumerate(reader):
            cells = [
                Cell(text=value, row_index=row_idx, col_index=col_idx)
                for col_idx, value in enumerate(row)
            ]
            lines.append(Line(cells=cells, line_number=row_idx + 1))
            if row_idx == 0:
                headers = list(row)
    return Document(lines=lines, source_path=path, source_format="csv", headers=headers)


def parse_xlsx(path: Path) -> Document:
    import openpyxl

    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as e:
        if "password" in str(e).lower() or "encrypted" in str(e).lower():
            raise ValueError(f"Password-protected files are not supported: {path.name}") from e
        raise
    lines: list[Line] = []
    headers: list[str] | None = None
    line_number = 0
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
            line_number += 1
            cells = [
                Cell(
                    text=str(value) if value is not None else "",
                    sheet_name=sheet_name,
                    row_index=row_idx,
                    col_index=col_idx,
                )
                for col_idx, value in enumerate(row)
            ]
            lines.append(Line(cells=cells, line_number=line_number))
            if row_idx == 0 and headers is None:
                headers = [str(v) if v is not None else "" for v in row]
    wb.close()
    return Document(lines=lines, source_path=path, source_format="xlsx", headers=headers)


def parse_docx(path: Path) -> Document:
    import docx

    doc = docx.Document(path)
    lines = [
        Line(cells=[Cell(text=para.text)], line_number=i + 1)
        for i, para in enumerate(doc.paragraphs)
    ]
    return Document(lines=lines, source_path=path, source_format="docx")


_PARSERS = {
    ".txt": parse_txt,
    ".csv": parse_csv,
    ".xlsx": parse_xlsx,
    ".docx": parse_docx,
}


def parse_file(path: Path) -> Document:
    ext = path.suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file format: {ext}")
    return parser(path)
